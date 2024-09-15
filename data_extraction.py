"""
import os
import PyPDF2
import docx
from cryptography.fernet import Fernet

# Generate a key for encryption and decryption
# You should save this key and use the same one for decryption
key = Fernet.generate_key()
cipher = Fernet(key)

def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfFileReader(file)
        for page_num in range(reader.numPages):
            page = reader.getPage(page_num)
            text += page.extractText()
    return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def encrypt_text(text):
    return cipher.encrypt(text.encode()).decode()

def decrypt_text(encrypted_text):
    return cipher.decrypt(encrypted_text.encode()).decode()

def append_text_to_file(text, output_file='user_data.txt'):
    encrypted_text = encrypt_text(text)
    with open(output_file, 'a', encoding='utf-8') as file:
        file.write(encrypted_text + "\n")

def read_and_decrypt_file(file_path='user_data.txt'):
    with open(file_path, 'r', encoding='utf-8') as file:
        encrypted_text = file.read().strip()
        return decrypt_text(encrypted_text)

def main():
    choice = input("Do you want to (1) Extract and append data or (2) Read and decrypt data? Enter 1 or 2: ")

    if choice == '1':
        file_path = input("Enter the path of the file: ")

        if not os.path.isfile(file_path):
            print("File not found.")
            return

        _, file_extension = os.path.splitext(file_path)

        if file_extension == '.txt':
            text = extract_text_from_txt(file_path)
        elif file_extension == '.pdf':
            text = extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            text = extract_text_from_docx(file_path)
        else:
            print("Unsupported file type.")
            return

        append_text_to_file(text)
        print("Text extracted, encrypted, and appended to user_data.txt")

    elif choice == '2':
        decrypted_text = read_and_decrypt_file()
        print("Decrypted text from user_data.txt:")
        print(decrypted_text)

    else:
        print("Invalid choice.")

if __name__ == "__main__":
    main()


import os
import PyPDF2
import docx
from cryptography.fernet import Fernet
import mammoth

def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text += page.extract_text()
    return text

def extract_text_from_docx(file_path):
    with open(file_path, "rb") as docx_file:
        result = mammoth.extract_raw_text(docx_file)
        text = result.value
        text = text.replace(" @", "@")
        return text
    # doc = docx.Document(file_path)
    # return "\n".join([para.text for para in doc.paragraphs])

def encrypt_text(text, cipher):
    return cipher.encrypt(text.encode()).decode()

def decrypt_text(encrypted_text, cipher):
    return cipher.decrypt(encrypted_text.encode()).decode()

def append_text_to_file(text, cipher, output_file='user_data.txt'):
    encrypted_text = encrypt_text(text, cipher)
    with open(output_file, 'a', encoding='utf-8') as file:
        file.write(encrypted_text + "\n")

def read_and_decrypt_file(cipher, file_path='user_data.txt'):
    with open(file_path, 'r', encoding='utf-8') as file:
        encrypted_texts = file.readlines()
        decrypted_texts = [decrypt_text(encrypted_text.strip(), cipher) for encrypted_text in encrypted_texts]
    return "\n".join(decrypted_texts)

def save_key(key, key_file='secret.key'):
    with open(key_file, 'wb') as file:
        file.write(key)

def load_key(key_file='secret.key'):
    if not os.path.exists(key_file):
        key = Fernet.generate_key()
        save_key(key, key_file)
    else:
        with open(key_file, 'rb') as file:
            key = file.read()
    return key

def main():
    key = load_key()
    cipher = Fernet(key)

    choice = input("Do you want to (1) Extract and append data or (2) Read and decrypt data? Enter 1 or 2: ")

    if choice == '1':
        file_path = input("Enter the path of the file: ")

        if not os.path.isfile(file_path):
            print("File not found.")
            return

        _, file_extension = os.path.splitext(file_path)

        if file_extension == '.txt':
            text = extract_text_from_txt(file_path)
        elif file_extension == '.pdf':
            text = extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            text = extract_text_from_docx(file_path)
        else:
            print("Unsupported file type.")
            return

        append_text_to_file(text, cipher)
        print("Text extracted, encrypted, and appended to user_data.txt")

    elif choice == '2':
        decrypted_text = read_and_decrypt_file(cipher)
        print("Decrypted text from user_data.txt:")
        print(decrypted_text)

    else:
        print("Invalid choice.")

if __name__ == "__main__":
    main()
"""

import os
import PyPDF2
import docx
from cryptography.fernet import Fernet
import mammoth

def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text += page.extract_text()
    return text

def extract_text_from_docx(file_path):
    with open(file_path, "rb") as docx_file:
        result = mammoth.extract_raw_text(docx_file)
        text = result.value
        text = text.replace(" @", "@")
        return text

def encrypt_text(text, cipher):
    return cipher.encrypt(text.encode()).decode()

def decrypt_text(encrypted_text, cipher):
    return cipher.decrypt(encrypted_text.encode()).decode()

def append_text_to_file(text, cipher, output_file):
    encrypted_text = encrypt_text(text, cipher)
    with open(output_file, 'a', encoding='utf-8') as file:
        file.write(encrypted_text + "\n")

def read_and_decrypt_file(cipher, file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        encrypted_texts = file.readlines()
        decrypted_texts = [decrypt_text(encrypted_text.strip(), cipher) for encrypted_text in encrypted_texts]
    return "\n".join(decrypted_texts)

def save_key(key, key_file):
    with open(key_file, 'wb') as file:
        file.write(key)

def load_key(key_file):
    if not os.path.exists(key_file):
        key = Fernet.generate_key()
        save_key(key, key_file)
    else:
        with open(key_file, 'rb') as file:
            key = file.read()
    return key

def log_processed_document(file_name):
    with open('docs_info.txt', 'a', encoding='utf-8') as file:
        file.write(file_name + "\n")

def main():
    choice = input("Do you want to (1) Extract and append data or (2) Read and decrypt data? Enter 1 or 2: ")

    if choice == '1':
        file_path = input("Enter the path of the file: ")

        if not os.path.isfile(file_path):
            print("File not found.")
            return

        file_name = os.path.splitext(os.path.basename(file_path))[0]
        output_file = f"{file_name}_data.txt"
        key_file = f"{file_name}_secret.key"

        key = load_key(key_file)
        cipher = Fernet(key)

        _, file_extension = os.path.splitext(file_path)

        if file_extension == '.txt':
            text = extract_text_from_txt(file_path)
        elif file_extension == '.pdf':
            text = extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            text = extract_text_from_docx(file_path)
        else:
            print("Unsupported file type.")
            return

        append_text_to_file(text, cipher, output_file)
        log_processed_document(file_name)
        print(f"Text extracted, encrypted, and saved to {output_file}")

    elif choice == '2':
        file_name = input("Enter the base name of the file (without extension): ")
        output_file = f"{file_name}_data.txt"
        key_file = f"{file_name}_secret.key"

        if not os.path.exists(output_file) or not os.path.exists(key_file):
            print("File or secret key not found.")
            return

        key = load_key(key_file)
        cipher = Fernet(key)
        decrypted_text = read_and_decrypt_file(cipher, output_file)
        print(f"Decrypted text from {output_file}:")
        print(decrypted_text)

    else:
        print("Invalid choice.")

if __name__ == "__main__":
    main()


